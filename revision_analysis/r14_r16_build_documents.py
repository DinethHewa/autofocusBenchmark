#!/usr/bin/env python3
"""Build manuscript, supplement, response package, tables, and rendered QC."""
from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path

import fitz
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "03_configs/revised_manuscript.md"
CLEAN = ROOT / "09_manuscript/clean/revised_manuscript_clean.docx"
HIGHLIGHT = ROOT / "09_manuscript/highlighted/revised_manuscript_highlighted.docx"
PDF = ROOT / "09_manuscript/rendered_qc/revised_manuscript.pdf"
PAGES = ROOT / "09_manuscript/rendered_qc/page_images"
RESP = ROOT / "10_response_package"
SUBMISSION = ROOT / "11_submission_package"
TABLE_MAIN = ROOT / "08_tables/main"
TABLE_SUPP = ROOT / "08_tables/supplementary"
TABLE_MACHINE = ROOT / "08_tables/machine_readable"

pdfmetrics.registerFont(TTFont("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
pdfmetrics.registerFont(TTFont("DejaVu-Bold", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"))

FIGURES = {
    "[[FIGURES1]]": (ROOT / "07_figures/supplementary/figureS1_entropy_correction.png", "Figure S1. Submitted and corrected TBF Histogram Entropy on the deterministic first stack."),
    "[[FIGURE2]]": (ROOT / "07_figures/main/figure2_reference_disagreement_example.png", "Figure 2. Deterministically selected largest WBC REF-B–REF-C disagreement."),
    "[[FIGURE3]]": (ROOT / "07_figures/main/figure3_reference_agreement.png", "Figure 3. Exact and within-one-slice REF-B–REF-C agreement by domain."),
    "[[FIGURE4]]": (ROOT / "07_figures/main/figure4_corrected_runtime_scaling.png", "Figure 4. Corrected repeated operator-kernel timing across resolutions."),
    "[[FIGURE5]]": (ROOT / "07_figures/main/figure5_rank_uncertainty.png", "Figure 5. Primary-score bootstrap intervals and conditional top-five frequencies."),
    "[[FIGURE6]]": (ROOT / "07_figures/main/figure6_domain_localization.png", "Figure 6. Domain-specific diagnostic-consensus displacement for the primary top six."),
    "[[FIGURES2]]": (ROOT / "07_figures/supplementary/figureS2_resampling_mechanism.png", "Figure S2. Controlled Roberts/Brenner response to scale, interpolation, and distortion."),
}


def tables() -> dict[str, tuple[pd.DataFrame, str]]:
    metadata = pd.read_csv(ROOT / "06_analysis_outputs/axial_units/acquisition_metadata_verified.csv")
    t1 = metadata[["domain", "patient_or_specimen_count", "slices_min", "slices_max", "image_dimensions", "axial_step_um", "objective_magnification", "numerical_aperture", "pixel_size_um", "licence"]].copy()
    t1.columns = ["Domain", "Patients/specimens", "Slices min", "Slices max", "Dimensions", "Step (µm)", "Objective", "NA", "Pixel size (µm)", "Licence"]
    ranking = pd.read_csv(ROOT / "06_analysis_outputs/corrected_scoring/corrected_final_rankings.csv").sort_values("rank")
    t2 = ranking.head(10)[["rank", "operator", "mu", "sigma", "G"]].copy(); t2.columns = ["Rank", "Operator", "μ", "σ", "G"]
    runtime = pd.read_csv(ROOT / "06_analysis_outputs/corrected_runtime/runtime_macro_micro_comparison.csv")
    top6 = ranking.head(6).operator.tolist(); t3 = runtime[runtime.operator.isin(top6)][["operator", "equal_domain_macro_median_kernel_ms", "slice_weighted_micro_median_kernel_ms", "equal_domain_macro_median_combined_ms"]].copy(); t3.columns = ["Operator", "Macro kernel (ms)", "Micro kernel (ms)", "Macro combined (ms)"]
    return {"[[TABLE1]]": (t1, "Table 1. Verified acquisition metadata; unreported fields are not inferred."), "[[TABLE2]]": (t2, "Table 2. Corrected primary top-ten ranking; G is candidate-pool-relative."), "[[TABLE3]]": (t3, "Table 3. Corrected native timing for the primary top six.")}


TABLES = tables()


def parse_markdown() -> list[dict]:
    blocks = []
    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line: continue
        if line in TABLES: blocks.append({"type": "table", "token": line})
        elif line in FIGURES: blocks.append({"type": "figure", "token": line})
        elif line.startswith("# "): blocks.append({"type": "title", "text": line[2:]})
        elif line.startswith("## "): blocks.append({"type": "h1", "text": line[3:]})
        elif line.startswith("### "): blocks.append({"type": "h2", "text": line[4:]})
        else: blocks.append({"type": "p", "text": line})
    return blocks


BLOCKS = parse_markdown()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]; section.top_margin = Inches(.65); section.bottom_margin = Inches(.65); section.left_margin = Inches(.72); section.right_margin = Inches(.72)
    normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(9); normal.paragraph_format.space_after = Pt(4); normal.paragraph_format.line_spacing = 1.05
    for style_name, size, color in (("Title", 16, "17365D"), ("Heading 1", 13, "17365D"), ("Heading 2", 11, "24527A")):
        style = doc.styles[style_name]; style.font.name = "Arial"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color); style.font.bold = True


def add_highlighted_text(paragraph, text: str, highlighted: bool, bold: bool = False) -> None:
    run = paragraph.add_run(text); run.bold = bold
    if highlighted: run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_docx_table(doc: Document, frame: pd.DataFrame, caption: str, highlighted: bool) -> None:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_highlighted_text(p, caption, highlighted, bold=True)
    table = doc.add_table(rows=1, cols=len(frame.columns)); table.style = "Table Grid"
    for index, column in enumerate(frame.columns):
        cell = table.rows[0].cells[index]; cell.text = str(column); shade_cell(cell, "D9EAF7")
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if isinstance(value, float): value = f"{value:.4f}"
            cells[index].text = str(value); cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"; run.font.size = Pt(6.5)
                    if highlighted: run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph()


def build_manuscript(path: Path, highlighted: bool) -> None:
    doc = Document(); configure_doc(doc)
    for block in BLOCKS:
        kind = block["type"]
        if kind == "title":
            p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_highlighted_text(p, block["text"], False)
        elif kind == "h1":
            p = doc.add_paragraph(style="Heading 1"); add_highlighted_text(p, block["text"], highlighted)
        elif kind == "h2":
            p = doc.add_paragraph(style="Heading 2"); add_highlighted_text(p, block["text"], highlighted)
        elif kind == "p":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            is_reference = any(previous.get("text") == "References" for previous in BLOCKS[max(0, BLOCKS.index(block)-1):BLOCKS.index(block)])
            add_highlighted_text(p, block["text"], highlighted and not re.match(r"^\d+\. ", block["text"]))
        elif kind == "table":
            add_docx_table(doc, *TABLES[block["token"]], highlighted)
        elif kind == "figure":
            image_path, caption = FIGURES[block["token"]]
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(image_path), width=Inches(6.35))
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_highlighted_text(cp, caption, highlighted, bold=True)
    path.parent.mkdir(parents=True, exist_ok=True); doc.save(path)


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("title", parent=base["Title"], fontName="DejaVu-Bold", fontSize=16, leading=19, textColor=colors.HexColor("#17365D"), alignment=TA_CENTER, spaceAfter=10),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontName="DejaVu-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#17365D"), spaceBefore=8, spaceAfter=4),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="DejaVu-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor("#24527A"), spaceBefore=6, spaceAfter=3),
        "p": ParagraphStyle("p", parent=base["BodyText"], fontName="DejaVu", fontSize=8.2, leading=10, alignment=TA_JUSTIFY, spaceAfter=4),
        "caption": ParagraphStyle("caption", parent=base["BodyText"], fontName="DejaVu-Bold", fontSize=7.5, leading=9, alignment=TA_CENTER, spaceAfter=6),
    }


def page_decorator(canvas, doc) -> None:
    canvas.saveState(); width, height = A4
    canvas.setFont("DejaVu", 6); canvas.setFillColor(colors.HexColor("#777777"))
    for line in range(5, 71, 5):
        y = height - 46 - (line - 1) * 10
        if y > 35: canvas.drawRightString(31, y, str(line))
    canvas.drawCentredString(width / 2, 20, f"Revised manuscript — page {doc.page}")
    canvas.restoreState()


def reportlab_table(frame: pd.DataFrame) -> Table:
    data = [list(frame.columns)] + [[f"{value:.4f}" if isinstance(value, float) else str(value) for value in row] for row in frame.itertuples(index=False, name=None)]
    available = A4[0] - 70; widths = [available / len(frame.columns)] * len(frame.columns)
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#D9EAF7")), ("GRID", (0,0), (-1,-1), .3, colors.HexColor("#777777")), ("FONTNAME", (0,0), (-1,0), "DejaVu-Bold"), ("FONTNAME", (0,1), (-1,-1), "DejaVu"), ("FONTSIZE", (0,0), (-1,-1), 5.4), ("VALIGN", (0,0), (-1,-1), "MIDDLE")]))
    return table


def build_pdf(path: Path) -> None:
    styles = pdf_styles(); story = []
    for block in BLOCKS:
        kind = block["type"]
        if kind in ("title", "h1", "h2", "p"):
            story.append(Paragraph(block["text"].replace("&", "&amp;"), styles[kind]))
        elif kind == "table":
            frame, caption = TABLES[block["token"]]; story.extend([Paragraph(caption, styles["caption"]), reportlab_table(frame), Spacer(1, 8)])
        else:
            image_path, caption = FIGURES[block["token"]]; image = Image(str(image_path)); ratio = min((A4[0]-78)/image.imageWidth, 5.0*inch/image.imageHeight); image.drawWidth=image.imageWidth*ratio; image.drawHeight=image.imageHeight*ratio; story.extend([image, Paragraph(caption, styles["caption"])])
    path.parent.mkdir(parents=True, exist_ok=True)
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=38, leftMargin=38, topMargin=40, bottomMargin=32, title="Revised manuscript jimaging-4524210").build(story, onFirstPage=page_decorator, onLaterPages=page_decorator)


def extract_docx_text(path: Path) -> list[str]:
    doc = Document(path); values = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            values.append("\t".join(cell.text for cell in row.cells))
    return [value for value in values if value.strip()]


def make_page_map() -> dict[str, dict]:
    document = fitz.open(PDF); rows = []; mapping = {}
    headings = [block["text"] for block in BLOCKS if block["type"] in ("h1", "h2")]
    for heading in headings:
        found = False
        for page_index, page in enumerate(document):
            hits = page.search_for(heading)
            if hits:
                line = max(1, int(round((hits[0].y0 - 40) / 10)) + 1)
                mapping[heading] = {"page": page_index + 1, "line": line}; rows.append({"section": heading, "page": page_index + 1, "line": line}); found = True; break
        if not found: rows.append({"section": heading, "page": "not found", "line": "not found"})
    pd.DataFrame(rows).to_csv(ROOT / "09_manuscript/rendered_qc/manuscript_page_line_map.csv", index=False)
    return mapping


def render_pages() -> None:
    PAGES.mkdir(parents=True, exist_ok=True); document = fitz.open(PDF)
    for index, page in enumerate(document):
        pix = page.get_pixmap(matrix=fitz.Matrix(1.35, 1.35), alpha=False); pix.save(PAGES / f"page_{index+1:03d}.png")


def export_tables() -> None:
    for directory in (TABLE_MAIN, TABLE_SUPP, TABLE_MACHINE): directory.mkdir(parents=True, exist_ok=True)
    for index, (token, (frame, caption)) in enumerate(TABLES.items(), 1):
        frame.to_csv(TABLE_MAIN / f"Table{index}.csv", index=False); frame.to_excel(TABLE_MAIN / f"Table{index}.xlsx", index=False)
    copies = {
        "STable1_all_32_corrected_rankings.csv": ROOT / "06_analysis_outputs/corrected_scoring/corrected_final_rankings.csv",
        "STable2_reference_agreement.csv": ROOT / "06_analysis_outputs/reference_audit/reference_agreement_by_domain.csv",
        "STable3_bootstrap_rank_frequencies.csv": ROOT / "06_analysis_outputs/statistical_inference/rank_frequencies.csv",
        "STable4_weight_sensitivity.csv": ROOT / "06_analysis_outputs/weight_sensitivity/deterministic_weight_and_alpha_sensitivity.csv",
        "STable5_resampling_rank_shifts.csv": ROOT / "06_analysis_outputs/resampling/resampling_rank_shifts.csv",
        "STable6_symbolic_provenance.csv": ROOT / "06_analysis_outputs/symbolic_audit/retained_composite_provenance.csv",
        "STable7_entropy_domain_summary.csv": ROOT / "06_analysis_outputs/corrected_entropy/submitted_vs_corrected_domain_summary.csv",
        "STable8_runtime_all.csv": ROOT / "06_analysis_outputs/corrected_runtime/corrected_runtime_per_measure_domain.csv",
        "STable9_axial_localization.csv": ROOT / "06_analysis_outputs/axial_units/operator_localization_axial_summary.csv",
    }
    for name, source in copies.items(): shutil.copy2(source, TABLE_SUPP / name)
    for source in sorted((ROOT / "06_analysis_outputs/raw_supplement").glob("*")): shutil.copy2(source, TABLE_MACHINE / source.name)


def build_supplement() -> Path:
    path = ROOT / "11_submission_package/revised_supplementary_material.docx"; doc = Document(); configure_doc(doc)
    doc.add_heading("Supplementary Material: Corrected Cross-Domain Focus-Measure Benchmark", 0)
    doc.add_paragraph("This supplement is generated from the corrected evidence. Machine-readable tables are authoritative where rounding differs.")
    sections = [
        ("S1. Immutable audit", "The input manifest hashes 5,128 pre-existing repository files and freezes the exact 32-operator pool and five stack counts. The end-of-revision verification compares those hashes without modifying the inputs."),
        ("S2. Formal criteria", "Absolute peak error is |p−y|. Range is the contiguous width at or above 95% of peak. False maxima exclude the selected global maximum. FWHM spans the first and last samples at or above half maximum. Noise is mean squared second difference. RRMSE is root mean squared clean/noisy curve difference divided by clean curve RMS. Runtime is corrected operator-kernel time unless explicitly labelled otherwise."),
        ("S3. Reference ladder", "REF-A is submitted historical replication; REF-B is fixed ten-voter diagnostic; REF-C is fixed four-voter disjoint confirmatory for 28 operators; REF-D is unavailable; REF-E awaits genuine experts."),
        ("S4. Statistical resampling", "Bootstrap replicates use 214 official WBC slide clusters and stacks for the other four domains. The release does not map WBC slides to its 72 patients. Criteria remain fixed policy components. Outputs include paired intervals, rank frequencies, family frequencies, leave-one-domain-out ranks, sigma sensitivity, reference sensitivity, and a WBC slide-cluster versus stack-level comparison."),
        ("S5. Controlled resampling", "Fourteen conditions separate native cached curves, aspect-preserving 0.5x/base/2x transformations, four interpolators, and square distortion. The 15-stack experiment is descriptive."),
        ("S6. Symbolic audit", "The 14 retained expressions are traced without rerunning GP. Exact and within-fold functional deduplication rules, scores, fold/seed origins, sizes, depths, and decisions are machine readable."),
        ("S7. Expert package", "The complete blinded annotation interface is supplied. No expert results appear because no genuine annotations were provided."),
    ]
    for heading, text in sections: doc.add_heading(heading, level=1); doc.add_paragraph(text)
    for index, (_, (frame, caption)) in enumerate(TABLES.items(), 1): add_docx_table(doc, frame, f"S{index}. {caption}", False)
    doc.add_heading("S8. Machine-readable inventory", level=1)
    for file in sorted(TABLE_SUPP.glob("*.csv")): doc.add_paragraph(file.name)
    path.parent.mkdir(parents=True, exist_ok=True); doc.save(path); return path


def claims_and_responses(page_map: dict[str, dict]) -> tuple[pd.DataFrame, pd.DataFrame]:
    claims = [
        ("32 operators and 26,100 stacks", "supported unchanged", "00_audit/dataset_inventory.csv"),
        ("Histogram Entropy valid on TBF", "supported after correction", "06_analysis_outputs/corrected_entropy/submitted_vs_corrected_domain_summary.csv"),
        ("ten-voter reference is independent", "removed", "06_analysis_outputs/reference_audit/reference_claim_limits.md"),
        ("gradient operators form the leading tier", "supported after correction", "06_analysis_outputs/corrected_scoring/corrected_final_rankings.csv"),
        ("Variance of Gradient is universally best", "weakened", "06_analysis_outputs/weight_sensitivity/dirichlet_rank_frequencies.csv"),
        ("localization is ground-truth error", "removed", "06_analysis_outputs/reference_audit/reference_claim_limits.md"),
        ("localization is algorithmic-consensus deviation", "supported after correction", "06_analysis_outputs/axial_units/operator_localization_axial_summary.csv"),
        ("10.7–24.4 ms per slice", "removed", "06_analysis_outputs/corrected_runtime/submitted_vs_corrected_runtime.csv"),
        ("leading kernels are computationally feasible", "supported after correction", "06_analysis_outputs/corrected_runtime/runtime_macro_micro_comparison.csv"),
        ("closed-loop real-time autofocus demonstrated", "removed", "06_analysis_outputs/corrected_runtime/RUNTIME_CORRECTION_FINDINGS.md"),
        ("ordinary 50-cell inference is confirmatory", "removed", "06_analysis_outputs/statistical_inference/STATISTICAL_FINDINGS.md"),
        ("equal-domain aggregation avoids WBC numerical dominance", "supported after correction", "06_analysis_outputs/domain_imbalance/aggregation_estimands.csv"),
        ("forced 1024 square identifies resampling mechanism", "removed", "06_analysis_outputs/resampling/RESAMPLING_FINDINGS.md"),
        ("14 composites have deterministic provenance", "supported after correction", "06_analysis_outputs/symbolic_audit/retained_composite_provenance.csv"),
        ("expert-independent accuracy", "blocked pending expert annotation", "06_analysis_outputs/expert_audit/EXPERT_AUDIT_BLOCKER.md"),
    ]
    claim_df = pd.DataFrame(claims, columns=["submitted_claim", "revision_classification", "evidence_artifact"]); claim_df.to_csv(RESP / "claim_evidence_matrix.csv", index=False)
    issues = [
        {
            "item": "R1.1", "reviewer": "Reviewer 1",
            "comment": "The manuscript states that the reference focus is constructed by averaging normalized focus curves from ten selected operators. Importantly, when an evaluated operator belongs to the voter set, that operator is removed from the consensus before the reference index is calculated. This avoids direct self-voting, which is sensible, but it appears to create another problem: a voter operator is evaluated against a nine-operator reference, whereas an operator outside the voter set is evaluated against the full ten-operator reference. Furthermore, each voter can potentially have a slightly different reference index. Consequently, the reported localization errors are not necessarily computed relative to one identical reference for every method?",
            "correct": "We agree.",
            "response": "We replaced the rotating comparison as the principal reference analysis. REF-A preserves the submitted construction only as historical evidence; REF-B is one fixed ten-voter diagnostic reference for all 32 operators; REF-C is one fixed four-voter disjoint reference for the 28 non-voters.",
            "result": "For WBC, voter-specific REF-A leave-one-out versus full-reference exact agreement ranged from 96.0% to 99.9%; REF-B versus REF-C exact agreement was 49.4%.",
            "section": "2.4. Reference ladder", "artifact": "Figures 2–3; reference_audit/reference_agreement_by_domain.csv", "limitation": "REF-B is comparable but self-including for its voters; REF-C is disjoint only for 28 operators. Neither is optical ground truth.",
        },
        {
            "item": "R1.2", "reviewer": "Reviewer 1",
            "comment": "The Friedman analysis treats each (domain, criterion) combination as one block, producing 50 blocks for the overall analysis. Bootstrap confidence intervals similarly resample the domain-criterion evaluation blocks, while the manuscript acknowledges that these blocks are correlated within domains. This is problematic because the ten criteria computed from a given domain are not independent replicates. Several are mathematically related properties of the same focus curves. Treating them as 50 effectively independent observations may substantially exaggerate the amount of independent evidence and produce overly optimistic p-values or confidence intervals.",
            "correct": "We agree.",
            "response": "The 50 domain–criterion cells are no longer used for confirmatory inference. We recalculate localization, domain scores, μ, σ, G, and ranks inside 1,000 clustered bootstrap replicates. The official WBC slide map was recovered and exactly reconciled to the cache, allowing 214-slide clustering; other domains use stacks.",
            "result": "Variance of Gradient was rank one in 100% of slide-clustered replicates. The leading top-five probabilities matched the stack-level WBC sensitivity to three decimals.",
            "section": "2.6. Statistical inference and domain imbalance", "artifact": "statistical_inference/rank_frequencies.csv; wbc_slide_cluster_vs_stack_bootstrap.csv", "limitation": "The release does not map its 214 slides to 72 patients, so residual dependence among slides from one patient cannot be modeled.",
        },
        {
            "item": "R1.3", "reviewer": "Reviewer 1",
            "comment": "The authors correctly note that WBC contributes approximately 98.7% of the 26,100 stacks and therefore use equal-domain aggregation as the primary analysis. This prevents WBC from numerically determining the aggregate score. However, equal-domain weighting creates the opposite issue: a domain containing only 30 stacks contributes exactly the same weight as the domain containing 25,773 stacks.",
            "correct": "We agree that these are different estimands rather than a single universally correct weighting.",
            "response": "We now report domain-specific results, equal-domain macro aggregation, per-stack micro aggregation, 1,000 balanced 30-per-domain subsamples, and leave-one-domain-out sensitivity separately.",
            "result": "WBC contributes 98.7% of stacks and therefore dominates the micro estimand; the gradient-family top tier persists across the reported estimands.",
            "section": "3.4. Resampling uncertainty, domain imbalance, and scoring policy", "artifact": "domain_imbalance/aggregation_estimands.csv", "limitation": "Equal-domain weighting protects against numerical dominance but does not equalize precision or define clinical domain importance.",
        },
        {
            "item": "R1.4", "reviewer": "Reviewer 1",
            "comment": "Within every domain-criterion cell, values are normalized to [0,1] across candidate measures. The authors correctly acknowledge that changing the candidate pool changes the normalized value scores. This means that the numerical value of G is not an absolute performance quantity and can change simply by adding or removing a poor or strong competitor. The manuscript should emphasize this more prominently, particularly in the Results and Conclusions. Please provide raw or physically interpretable values for the major criteria for all operators in the supplementary material so readers can reproduce comparisons without dependence on the chosen pool.",
            "correct": "We agree.",
            "response": "The Abstract, Methods, Results, table captions, Discussion, and Conclusion now call G a candidate-pool-relative decision score. Raw and normalized values for every operator, domain, and criterion are exported in CSV/XLSX, together with per-stack localization.",
            "result": "The supplement contains all 32 operators × 5 domains × 10 criteria and a per-stack localization table.",
            "section": "2.5. Criteria and candidate-pool-relative score", "artifact": "raw_supplement/all_operators_domains_criteria_long.csv/.xlsx", "limitation": "Rankings remain conditional on the frozen 32-operator pool and declared normalization policy.",
        },
        {
            "item": "R1.5", "reviewer": "Reviewer 1",
            "comment": "The leading gradient operators are reported at approximately 10.7–24.4 ms per slice, and the manuscript concludes that this makes them compatible with real-time autofocus. The runtime protocol, however, excludes image I/O and one-time setup and reports computation under a single CPU protocol. Furthermore, because WBC contributes nearly all slices and consists of 200 × 200 crops, an across-slice native-resolution average will be heavily dominated by this relatively small image format.",
            "correct": "We agree.",
            "response": "We retired the submitted timing range and repeated the experiment with matched preprocessing, two warm-ups, seven randomized repeats, component-wise I/O/preprocessing/kernel/combined timing, resolution-specific results, and explicit macro and micro estimands.",
            "result": "Top-six equal-domain macro native kernel medians were 11.1–29.1 ms; slice-weighted micro medians were 0.27–0.72 ms.",
            "section": "3.5. Corrected runtime", "artifact": "Table 3; Figure 4; corrected_runtime outputs", "limitation": "One workstation and 15 deterministic images were timed; exposure, stage motion, search, transfer, and controller overhead were not measured.",
        },
        {
            "item": "R1.6", "reviewer": "Reviewer 1",
            "comment": "After resampling every image to 1024 × 1024, the Spearman rank correlation falls to 0.418 and only three of the top five operators remain. This is an important finding and, in my view, one of the manuscript's more useful practical observations. However, resizing heterogeneous images to one square grid simultaneously changes sampling density, interpolation artifacts, and potentially effective spatial-frequency content. It can therefore be difficult to determine why a particular operator changes rank.",
            "correct": "We agree.",
            "response": "A 14-condition controlled study separates aspect-preserving scale changes, four interpolation methods, and square-grid distortion, using the same frozen operator pool on 15 deterministic stacks.",
            "result": "For 0.5× area downsampling, rank improvement correlated weakly with noise reduction (ρ=0.104), but more strongly with reduced false maxima (ρ=0.582) and narrower FWHM (ρ=0.526).",
            "section": "3.6. Controlled resampling and Roberts/Brenner behavior", "artifact": "Figure S2; resampling/roberts_brenner_mechanism.csv", "limitation": "The mechanism analysis is descriptive and based on 15 stacks.",
        },
        {
            "item": "R1.7", "reviewer": "Reviewer 1",
            "comment": "The LODO design is conceptually reasonable, and the manuscript appropriately labels symbolic fusion as exploratory. However, the Results state that 14 symbolic composites were retained without sufficiently clear explanation in the main Methods of precisely how these 14 were selected from the folds, seeds, and candidate expressions.",
            "correct": "We agree that the submitted description was insufficient.",
            "response": "The Methods now specify the final refit plus leave-one-domain-out representatives, exact-expression and within-fold functional deduplication, correlation threshold, deterministic ordering, and node/depth constraints. The full provenance is exported without rerunning GP.",
            "result": "All 14 retained composites are traced to expression, fold, seed, scores, size, depth, duplicate group, filtering decision, and retention reason.",
            "section": "2.9. Symbolic composites", "artifact": "symbolic_audit/retained_composite_provenance.csv", "limitation": "The symbolic component remains exploratory.",
        },
        {
            "item": "R1.8", "reviewer": "Reviewer 1",
            "comment": "While the paper evaluates several statistical, texture, and entropy-based focus measures (e.g., GLCM Contrast, Histogram Entropy, Normalized Variance), the discussion of how image complexity and spatial relationships influence feature extraction in digital pathology could be enriched. Smear microscopy and digital pathology slides present variable cellular distributions and spatial architectures that directly affect focus curve morphology and entropy profiles. To contextualize the benchmark within the broader scope of quantitative digital pathology and spatial/structural complexity, please cite and discuss: Li, X. (2024). ‘Deciphering cell to cell spatial relationship for pathology images using SpatialQPFs.’ Scientific Reports, 14, 29585; and Li, X., Ren, X., & Venugopal, R. (2025). ‘Entropy measures for quantifying complexity in digital pathology and spatial omics.’ iScience, 28(6).",
            "correct": "We agree that these references provide useful context when treated critically.",
            "response": "Both references were added to the Introduction and Discussion. We distinguish cell-coordinate spatial statistics and spatial entropy from our global gray-level Histogram Entropy and local GLCM Contrast, and explicitly state that the cited work does not validate an autofocus operator.",
            "result": "The new discussion explains how cellular density, stain distribution, and architecture can affect entropy and texture curves independently of defocus.",
            "section": "4. Discussion", "artifact": "References 29–30", "limitation": "No segmented-cell spatial descriptors were computed in this benchmark.",
        },
        {
            "item": "R2.1", "reviewer": "Reviewer 2",
            "comment": "The study employs multiple datasets, and certain metrics—such as focal plane selection error—are reported in units of ‘slices.’ The authors should clarify whether the acquisition parameters across these datasets are consistent. In particular, was the axial step size uniform across all z-stacks? If not, how might such discrepancies affect the calculation of the evaluation metrics?",
            "correct": "We agree that slice errors are not physically comparable without domain-specific acquisition metadata.",
            "response": "We verified axial steps from official records and report slice, micrometre, normalized axial, exact-match, and within-one-slice results by domain. No cross-domain result is multiplied by the WBC step.",
            "result": "Verified steps are WBC 0.4 µm, TBI 2.5 µm, and PBS/BMA/TBF 0.5 µm.",
            "section": "2.2. Datasets and acquisition metadata", "artifact": "Table 1; axial_units outputs", "limitation": "Pixel size and several other optical fields remain unreported.",
        },
        {
            "item": "R2.2", "reviewer": "Reviewer 2",
            "comment": "There appears to be a substantial imbalance in the number of samples across dataset categories—for instance, 25,773 WBC images versus only dozens to hundreds for other categories. Could this imbalance compromise the robustness of the reported performance metrics?",
            "correct": "Yes; it changes the estimand and the precision.",
            "response": "We separate domain-specific, equal-domain macro, per-stack micro, balanced repeated-subsample, and leave-one-domain-out results and explain the interpretation of each.",
            "result": "WBC is 98.7% of all stacks, so the micro result is WBC-dominated; the equal-domain result is not.",
            "section": "3.4. Resampling uncertainty, domain imbalance, and scoring policy", "artifact": "domain_imbalance outputs", "limitation": "Small-domain uncertainty and generalization beyond the five observed domains remain.",
        },
        {
            "item": "R2.3", "reviewer": "Reviewer 2",
            "comment": "There appears to be a substantial imbalance in the number of samples across dataset categories—for instance, 25,773 WBC images versus only dozens to hundreds for other categories. Could this imbalance compromise the robustness of the reported performance metrics?",
            "correct": "We address this repeated comment separately and cross-reference the full response to R2.2.",
            "response": "As described for R2.2, no single aggregate is presented as universal; all four estimands and leave-one-domain-out results are retained separately.",
            "result": "The family-level conclusion persists, while numerical scores and lower ranks vary across estimands.",
            "section": "3.4. Resampling uncertainty, domain imbalance, and scoring policy", "artifact": "domain_imbalance outputs", "limitation": "See R2.2.",
        },
        {
            "item": "R2.4", "reviewer": "Reviewer 2",
            "comment": "The authors report a focal plane selection time of 10.7–24.4 ms. However, in many modern imaging systems, exposure times can be reduced to just a few milliseconds. How do the authors interpret the impact of the discrepancy between data acquisition time and focus computation time on real-time autofocus performance? Moreover, the current study selects the optimal focal plane from an already acquired image sequence. In practical imaging scenarios, however, the goal is not to choose the best image from a known sequence, but to predict the position of optimal focus based on parameters from the current image. Therefore, the authors should more clearly define and discuss whether their proposed method is intended for the data acquisition stage of microscopy or for downstream data processing tasks—such as AI-based image analysis.",
            "correct": "We agree that the submitted scope was too broad.",
            "response": "The manuscript now defines the study as retrospective selection from acquired z-stacks. Focus measures may inform an acquisition-time search objective or downstream selection, but no acquisition controller, axial-position predictor, or closed-loop real-time system is claimed.",
            "result": "Runtime is reported as component-wise operator feasibility, with the obsolete real-time claim removed.",
            "section": "2.7. Corrected runtime protocol", "artifact": "Table 3; Figure 4; claim_evidence_matrix.csv", "limitation": "Acquisition exposure, motion, search, and stopping logic were not tested.",
        },
        {
            "item": "R2.5", "reviewer": "Reviewer 2",
            "comment": "While the paper evaluates focus assessment on publicly available datasets, it does not include a single example of a microscopic image. The authors should present representative images to illustrate the research process—for instance, demonstrating which operators select higher-quality focal planes—so that readers can gain a more intuitive understanding of the microscopic data involved.",
            "correct": "We agree.",
            "response": "Figures 1 and 2 add deterministic, non-cherry-picked representative sequences and a reference-disagreement example with focus curves. A licence manifest records all derivatives.",
            "result": "WBC, PBS, BMA, and TBF examples are included; TBI derivatives are omitted because the official licence was not reported.",
            "section": "3.2. Reference construction", "artifact": "Figures 1–2; 07_figures/licence_manifest.csv", "limitation": "TBI raw-image derivatives remain excluded pending licence confirmation.",
        },
        {
            "item": "R2-English", "reviewer": "Reviewer 2",
            "comment": "The English could be improved to more clearly express the research.",
            "correct": "We agree.",
            "response": "The manuscript was rewritten throughout for shorter sentences, explicit definitions, consistent reference terminology, corrected duplicated text and punctuation, and clearer separation of findings from limitations.",
            "result": "The clean and highlighted manuscripts contain identical scientific text and passed the document comparison check.",
            "section": "1. Introduction", "artifact": "clean/highlighted manuscripts; document_validation.json", "limitation": "Final journal copy-editing may still make house-style changes.",
        },
        {
            "item": "R3.1", "reviewer": "Reviewer 3",
            "comment": "The clinical or engineering significance of the conclusion that the mean error ranges from 0.11 to 0.19 focal slices in an actual autofocus system has not been sufficiently discussed.",
            "correct": "We agree.",
            "response": "The revision reports domain-specific slices and micrometres, exact and within-one-slice rates, and adds engineering context. The PBS/BMA/TBF source reports an approximately 0.5 µm depth of field for its 100×/1.4 NA green-light setup, equal to one acquired slice interval. We explicitly avoid treating this as a clinical tolerance.",
            "result": "For Variance of Gradient, mean REF-B deviations are 0.025–0.260 µm in PBS/BMA/TBF, below the source-reported ~0.5 µm depth-of-field scale.",
            "section": "4. Discussion", "artifact": "axial_units outputs; Reference 31", "limitation": "No expert acceptable-focus interval, stage-repeatability measurement, or downstream clinical tolerance is available; no DoF estimate is assigned to WBC or TBI.",
        },
        {
            "item": "R3.2", "reviewer": "Reviewer 3",
            "comment": "What is the basis for selecting α = 0.7 in Equation (5)? Is it an arbitrary preset? Although a sensitivity analysis for α (ranging from 0.5 to 0.9) has been conducted, it is recommended that the rationale for this design choice be stated in the main text.",
            "correct": "The concern is correct.",
            "response": "We identify α=0.7 as an author-defined policy value encoded in the submitted analysis and do not claim prespecification. Sensitivity covers α=0,0.1,…,1.0.",
            "result": "The primary winner is reported only under the declared policy; α=0 is labelled an extreme dispersion-only stress test.",
            "section": "2.5. Criteria and candidate-pool-relative score", "artifact": "weight_sensitivity outputs; Figure S3", "limitation": "No independent clinical utility function establishes an optimal α.",
        },
        {
            "item": "R3.3", "reviewer": "Reviewer 3",
            "comment": "Regarding the weight allocation in Table 3, it is recommended to clarify whether the weights were set by domain experts or obtained through data-driven optimization.",
            "correct": "We agree that provenance must be explicit.",
            "response": "The weights are author-defined constants from the submitted code; they were neither expert-elicited nor data-optimized in this revision, and repository history does not establish prespecification. Equal, localization-heavy, runtime, and 1,000 Dirichlet-weight sensitivities were added.",
            "result": "Variance of Gradient had 70.7% rank-one probability over Dirichlet draws, so no invariant individual winner is claimed.",
            "section": "2.5. Criteria and candidate-pool-relative score", "artifact": "weight_sensitivity/weight_provenance.md; Figure S3", "limitation": "Policy uncertainty remains.",
        },
        {
            "item": "R3.4", "reviewer": "Reviewer 3",
            "comment": "In Figure 4, after resampling, the Roberts Focus Measure jumps from rank 9 to rank 1, while the Brenner Gradient rises from rank 6 to rank 2—this dramatic change warrants separate discussion.",
            "correct": "We agree.",
            "response": "We added a controlled resampling experiment and a dedicated Results/Discussion analysis of scale, interpolation, peak position, false maxima, curve noise, FWHM, and pixel support.",
            "result": "The proposed noise-only mechanism was not supported (ρ=0.104 at 0.5× area); false-maxima and FWHM reductions showed stronger associations (ρ=0.582 and 0.526).",
            "section": "3.6. Controlled resampling and Roberts/Brenner behavior", "artifact": "Figure S2; resampling outputs", "limitation": "Associations are descriptive and do not prove mechanism.",
        },
        {
            "item": "R4.1", "reviewer": "Reviewer 4",
            "comment": "In this work, the central limitation of the benchmark is that the ‘reference focus’ is itself constructed from focus measures being evaluated, rather than from an independent optical or experimental ground truth. Although leave-one-operator-out exclusion removes direct self-voting, it does not remove family-level dependence. For example, when a gradient operator is evaluated, several highly correlated derivative-based operators remain in the principal consensus. Therefore, the very small reported localization errors (approximately 0.11–0.19 slices for the leading gradient measures) partly quantify agreement with a related ensemble rather than accuracy relative to the true optical best-focus plane. The independent non-derivative reconstruction is an important robustness check, but it does not fully resolve this issue. In particular, the supplementary results show only 49.5% exact reference-index agreement for WBC between the ten-voter and four-voter constructions, with a mean absolute shift of 0.95 slices and a 90th-percentile shift of 2 slices. These discrepancies are substantial relative to the reported 0.11–0.19-slice localization errors. The authors should therefore temper the interpretation of the localization values and, if possible, validate the consensus reference on an independent subset using hardware-calibrated axial positions or blinded expert assessment.",
            "correct": "We fully agree with the methodological limitation.",
            "response": "All accuracy and ground-truth wording was removed. Localization is called algorithmic-consensus deviation. REF-B is explicitly diagnostic, REF-C is disjoint only for 28 operators, and their complete disagreement distributions are reported. A blinded two-assessor-plus-adjudicator package is complete and now samples 30 distinct official WBC slides, but no genuine expert labels were available for this revision.",
            "result": "WBC REF-B–REF-C exact agreement is 49.4%, mean absolute shift 0.950 slices, P90 2 slices, and maximum 7 slices.",
            "section": "4. Discussion", "artifact": "Figures 2–3; reference_audit outputs; expert_audit package", "limitation": "Independent expert/hardware validation remains unrun and is not fabricated.",
        },
    ]
    rows = []
    for issue in issues:
        loc = page_map.get(issue["section"], {"page": "not found", "line": "not found"})
        rows.append({
            "item": issue["item"], "reviewer": issue["reviewer"], "reviewer_comment": issue["comment"],
            "whether_reviewer_is_correct": issue["correct"], "direct_response": issue["response"],
            "analysis_or_correction": issue["response"], "main_result": issue["result"],
            "manuscript_section": issue["section"], "page": loc["page"], "line": loc["line"],
            "artifact": issue["artifact"], "remaining_limitation": issue["limitation"],
        })
    response_df = pd.DataFrame(rows); response_df.to_csv(RESP / "reviewer_comment_matrix.csv", index=False)
    return claim_df, response_df


def generic_docx(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = Document(); configure_doc(doc); doc.add_heading(title, 0)
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs: doc.add_paragraph(paragraph)
    doc.save(path)


def generic_pdf(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    styles = pdf_styles(); story = [Paragraph(title, styles["title"])]
    for heading, paragraphs in sections:
        story.append(Paragraph(heading, styles["h1"])); story.extend(Paragraph(p.replace("&", "&amp;"), styles["p"]) for p in paragraphs)
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=36).build(story)


def build_response_documents(claim_df: pd.DataFrame, response_df: pd.DataFrame) -> None:
    summary = [
        "We thank the Academic Editor and four reviewers for their detailed and constructive assessment. We addressed all 18 numbered comments separately and also acknowledge Reviewer 2’s English-language assessment.",
        "First, our revision audit identified a dtype defect in Histogram Entropy: a fixed 0–256 histogram range excluded valid uint16 TBF intensities. We implemented a documented dtype-safe definition, passed 11 tests, recomputed all five domains, and regenerated all dependent evidence. All 182 TBF entropy curves and peaks changed; the corrected ten-voter REF-B reference remained unchanged through median-vote robustness, which does not excuse the original defect.",
        "Second, we replaced ambiguous reference labels with REF-A–REF-E, separated the fixed diagnostic REF-B from the disjoint confirmatory REF-C, demoted the 50-cell inferential analysis, recovered the official WBC image-to-slide map for 214-slide clustered bootstrap, rebuilt runtime measurement, controlled resampling factors, and narrowed every claim to what the evidence supports.",
        "Third, we added the two reviewer-suggested Li papers with critical—not ceremonial—discussion, added source-supported depth-of-field engineering context, representative microscopy images, raw machine-readable results, and a complete blinded expert package. Genuine expert results remain unavailable and are not fabricated.",
    ]
    sections = [("Summary of Substantive Changes", summary)]
    current_reviewer = None
    for row in response_df.itertuples(index=False):
        if row.reviewer != current_reviewer:
            current_reviewer = row.reviewer
            sections.append((current_reviewer, ["Each comment is reproduced below, followed by our response, action, result, location, evidence artifact, and remaining limitation."]))
        sections.append((f"{row.item}", [f"Reviewer comment: {row.reviewer_comment}", f"Assessment: {row.whether_reviewer_is_correct}", f"Response and action: {row.direct_response}", f"Main result: {row.main_result}", f"Location: {row.manuscript_section}, page {row.page}, line {row.line}.", f"Artifact: {row.artifact}.", f"Remaining limitation: {row.remaining_limitation}"]))
    generic_docx(RESP / "response_to_reviewers.docx", "Response to Reviewers — jimaging-4524210", sections)
    generic_pdf(RESP / "response_to_reviewers.pdf", "Response to Reviewers — jimaging-4524210", sections)
    cover_sections = [("Dear Academic Editor and Reviewers", ["We submit the major revision of “Cross-Domain Benchmarking of Focus Measures for Smear-Microscopy Autofocus Under a Consensus-Audited Reference” (jimaging-4524210). We thank the reviewers for identifying important concerns about reference comparability, inferential independence, domain imbalance, runtime interpretation, resampling, physical units, practical significance, and independent validation.", "Our audit also identified and transparently corrected a uint16 Histogram Entropy defect and an inadequate submitted runtime protocol. All five domains were recomputed where required, all dependent scores were regenerated, and the immutable submitted evidence was preserved. The response document addresses all 18 numbered comments separately and acknowledges Reviewer 2’s English-language assessment.", "The revised manuscript uses unambiguous REF-A–REF-E reference terminology, uses 214-slide clustered WBC bootstrap, reports domain/macro/micro/balanced estimands, adds verified physical units and source-supported depth-of-field context, critically discusses the two suggested Li papers, controls resampling factors, explains the 14 symbolic composites, includes representative images, and supplies a blinded expert package. The principal conclusion is deliberately narrower: gradient-family operators form a strong tier; Variance of Gradient is first only under the declared candidate pool and policy.", "No expert labels or missing patient identifiers were fabricated. Independent expert/hardware validation remains a stated limitation, and the complete blinded instrument is supplied for future execution."]), ("Sincerely", ["Dineth Hewavitharana, on behalf of the authors"])]
    generic_docx(RESP / "cover_letter.docx", "Major Revision Cover Letter", cover_sections); generic_pdf(RESP / "cover_letter.pdf", "Major Revision Cover Letter", cover_sections)


def main() -> int:
    for directory in (CLEAN.parent, HIGHLIGHT.parent, PDF.parent, RESP, SUBMISSION, TABLE_MAIN, TABLE_SUPP, TABLE_MACHINE): directory.mkdir(parents=True, exist_ok=True)
    export_tables(); build_manuscript(CLEAN, False); build_manuscript(HIGHLIGHT, True); build_pdf(PDF); render_pages(); page_map = make_page_map(); supplement = build_supplement()
    clean_text = extract_docx_text(CLEAN); highlighted_text = extract_docx_text(HIGHLIGHT)
    claims, responses = claims_and_responses(page_map); build_response_documents(claims, responses)
    change_rows = [{"section": heading, "change": "substantively revised from corrected evidence", "evidence": artifact} for heading, artifact in (("Abstract", "corrected scoring/reference/runtime"), ("Methods", "Levels 0–12 configurations and official WBC slide map"), ("Results", "corrected outputs and figures"), ("Discussion and limitations", "claim evidence matrix, engineering context, and reference limits"), ("Tables and figures", "08_tables and 07_figures"), ("References", "duplicate author names corrected; Li 2024 and Li et al. 2025 critically integrated"))]
    pd.DataFrame(change_rows).to_csv(RESP / "manuscript_change_log.csv", index=False)
    (RESP / "extension_request_draft.md").write_text("""# Optional extension request for blinded expert assessment

Dear Editorial Office,

We are completing the major revision of manuscript jimaging-4524210. All reviewer comments have been addressed in the revised manuscript and point-by-point response. Reviewer 4 suggested independent blinded expert assessment if possible. We have prepared and validated a complete two-assessor-plus-adjudicator instrument, including an unbiased domain-stratified sample and a separately labelled disagreement-enriched qualitative set. If the Editorial Office considers completed expert assessment essential for this revision round, we respectfully request a short extension to coordinate qualified assessors and report genuine annotations. We will not fabricate or substitute algorithmic labels for expert judgments.

Sincerely,  
The Authors
""", encoding="utf-8")
    validation = {"status": "PASS" if clean_text == highlighted_text and len(responses) == 19 else "FAIL", "clean_highlighted_scientific_text_identical": clean_text == highlighted_text, "clean_paragraph_table_records": len(clean_text), "highlighted_paragraph_table_records": len(highlighted_text), "manuscript_pdf_pages": len(fitz.open(PDF)), "rendered_page_images": len(list(PAGES.glob("page_*.png"))), "numbered_reviewer_comments_addressed": len(responses[responses.item.str.match(r"R[1-4]\.\d+")]), "english_language_assessment_addressed": bool((responses.item == "R2-English").any()), "response_items_total": len(responses), "reviewer_comments_source": "official decision letter supplied by author", "two_reviewer_suggested_Li_papers": "included and critically discussed", "supplement_exists": supplement.exists()}
    (ROOT / "09_manuscript/rendered_qc/document_validation.json").write_text(json.dumps(validation, indent=2) + "\n", encoding="utf-8")
    (ROOT / "12_logs/r14_r16_build_documents.log").write_text(json.dumps(validation, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(validation, indent=2)); return 0 if validation["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
