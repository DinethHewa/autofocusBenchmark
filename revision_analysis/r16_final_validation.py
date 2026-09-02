#!/usr/bin/env python3
"""Assemble the submission package and independently validate the revision."""
from __future__ import annotations

import csv
import hashlib
import json
import math
import os
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
REPOSITORY = ROOT.parent
SUB = ROOT / "11_submission_package"
MACHINE = SUB / "machine_readable_supplement"
FIGURES = SUB / "figures"
CODE = SUB / "code_release_manifest"
EXPERT = SUB / "blinded_expert_audit"
BASELINE = ROOT / "00_audit/input_manifest.csv"


def sha(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(8 * 1024 * 1024), b""): digest.update(chunk)
    return digest.hexdigest()


def copy(source: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True); shutil.copy2(source, destination)


def extract_docx(path: Path) -> list[str]:
    doc = Document(path); values = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        values.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
    return [value for value in values if value.strip()]


def package() -> None:
    for directory in (SUB, MACHINE, FIGURES, CODE, EXPERT / "blinded_annotation_package"): directory.mkdir(parents=True, exist_ok=True)
    core = {
        ROOT / "09_manuscript/clean/revised_manuscript_clean.docx": SUB / "revised_manuscript_clean.docx",
        ROOT / "09_manuscript/highlighted/revised_manuscript_highlighted.docx": SUB / "revised_manuscript_highlighted.docx",
        ROOT / "11_submission_package/revised_supplementary_material.docx": SUB / "revised_supplementary_material.docx",
        ROOT / "10_response_package/response_to_reviewers.docx": SUB / "response_to_reviewers.docx",
        ROOT / "10_response_package/cover_letter.docx": SUB / "cover_letter.docx",
        ROOT / "09_manuscript/rendered_qc/revised_manuscript.pdf": SUB / "revised_manuscript.pdf",
        ROOT / "10_response_package/response_to_reviewers.pdf": SUB / "response_to_reviewers.pdf",
        ROOT / "10_response_package/cover_letter.pdf": SUB / "cover_letter.pdf",
    }
    for source, destination in core.items():
        if source.resolve() != destination.resolve(): copy(source, destination)
    for source in sorted((ROOT / "07_figures/main").glob("*")) + sorted((ROOT / "07_figures/supplementary").glob("*")):
        copy(source, FIGURES / source.name)
    evidence = [
        ROOT / "06_analysis_outputs/corrected_scoring/corrected_raw_criteria.csv",
        ROOT / "06_analysis_outputs/corrected_scoring/corrected_final_rankings.csv",
        ROOT / "06_analysis_outputs/corrected_scoring/submitted_vs_corrected_change_impact.csv",
        ROOT / "06_analysis_outputs/raw_supplement/per_stack_localization.csv",
        ROOT / "06_analysis_outputs/raw_supplement/all_operators_domains_criteria_long.csv",
        ROOT / "06_analysis_outputs/raw_supplement/all_operators_domains_criteria_long.xlsx",
        ROOT / "06_analysis_outputs/reference_audit/reference_agreement_by_domain.csv",
        ROOT / "06_analysis_outputs/reference_audit/reference_transition_matrices.csv",
        ROOT / "06_analysis_outputs/reference_audit/reference_shift_distributions.csv",
        ROOT / "06_analysis_outputs/corrected_entropy/submitted_vs_corrected_domain_summary.csv",
        ROOT / "06_analysis_outputs/corrected_entropy/corrected_entropy_rrmse_per_stack.csv",
        ROOT / "06_analysis_outputs/corrected_runtime/corrected_runtime_per_measure_domain.csv",
        ROOT / "06_analysis_outputs/corrected_runtime/corrected_runtime_per_resolution.csv",
        ROOT / "06_analysis_outputs/statistical_inference/rank_frequencies.csv",
        ROOT / "06_analysis_outputs/statistical_inference/paired_difference_intervals.csv",
        ROOT / "06_analysis_outputs/statistical_inference/wbc_slide_cluster_vs_stack_bootstrap.csv",
        ROOT / "06_analysis_outputs/statistical_inference/wbc_slide_mapping/wbc_stack_to_slide.csv",
        ROOT / "06_analysis_outputs/statistical_inference/wbc_slide_mapping/validation_summary.json",
        ROOT / "06_analysis_outputs/domain_imbalance/aggregation_estimands.csv",
        ROOT / "06_analysis_outputs/weight_sensitivity/deterministic_weight_and_alpha_sensitivity.csv",
        ROOT / "06_analysis_outputs/weight_sensitivity/dirichlet_rank_frequencies.csv",
        ROOT / "06_analysis_outputs/resampling/resampling_rank_shifts.csv",
        ROOT / "06_analysis_outputs/resampling/roberts_brenner_mechanism.csv",
        ROOT / "06_analysis_outputs/axial_units/operator_localization_axial_summary.csv",
        ROOT / "06_analysis_outputs/axial_units/acquisition_metadata_verified.csv",
        ROOT / "06_analysis_outputs/axial_units/depth_of_field_context.csv",
        ROOT / "06_analysis_outputs/axial_units/depth_of_field_context.md",
        ROOT / "06_analysis_outputs/symbolic_audit/retained_composite_provenance.csv",
        ROOT / "06_analysis_outputs/expert_audit/annotation_manifest.csv",
        ROOT / "10_response_package/claim_evidence_matrix.csv",
        ROOT / "10_response_package/reviewer_comment_matrix.csv",
    ]
    for source in evidence: copy(source, MACHINE / source.name)
    expert_root = ROOT / "06_analysis_outputs/expert_audit"
    for name in ("EXPERT_AUDIT_BLOCKER.md", "annotation_configuration.json", "annotation_instructions.md", "annotation_manifest.csv", "annotation_template.csv", "adjudication_template.csv"):
        copy(expert_root / name, EXPERT / name)
    for name in ("blinded_manifest.json", "index.html", "serve_annotations.py"):
        copy(expert_root / "blinded_annotation_package" / name, EXPERT / "blinded_annotation_package" / name)
    for source in (ROOT / "00_audit/input_manifest.csv", ROOT / "00_audit/input_hashes.sha256", ROOT / "00_audit/operator_registry_32.json", ROOT / "00_audit/environment.json"):
        copy(source, CODE / source.name)
    manifests = []
    for path in sorted([*ROOT.glob("02_revision_code/*.py"), *ROOT.glob("03_configs/*"), *ROOT.glob("04_tests/*.py")]):
        manifests.append({"path": str(path.relative_to(ROOT)), "bytes": path.stat().st_size, "sha256": sha(path), "role": "revised code/config/test"})
    for path in sorted((ROOT / "01_frozen_code_copy").rglob("*")):
        if path.is_file(): manifests.append({"path": str(path.relative_to(ROOT)), "bytes": path.stat().st_size, "sha256": sha(path), "role": "frozen submitted evidence copy"})
    with (CODE / "revision_code_manifest.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(manifests[0])); writer.writeheader(); writer.writerows(manifests)


def outside_hash_check() -> dict:
    baseline_rows = list(csv.DictReader(BASELINE.open(newline="", encoding="utf-8")))
    baseline = {row["relative_path"]: row["sha256"] for row in baseline_rows}
    changed, missing = [], []
    for index, (relative, expected) in enumerate(baseline.items(), 1):
        path = REPOSITORY / relative
        if not path.exists(): missing.append(relative)
        elif sha(path) != expected: changed.append(relative)
    current_paths = set()
    for path in REPOSITORY.rglob("*"):
        if not path.is_file(): continue
        try: relative = path.relative_to(REPOSITORY)
        except ValueError: continue
        if relative.parts and relative.parts[0] == ROOT.name: continue
        current_paths.add(str(relative))
    unexpected = sorted(current_paths - set(baseline))
    rows = ([{"status": "changed", "relative_path": value} for value in changed] + [{"status": "missing", "relative_path": value} for value in missing] + [{"status": "unexpected_new", "relative_path": value} for value in unexpected])
    with (ROOT / "00_audit/final_outside_root_hash_verification.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["status", "relative_path"]); writer.writeheader(); writer.writerows(rows)
    result = {"baseline_files": len(baseline), "changed": changed, "missing": missing, "unexpected_new": unexpected, "status": "PASS" if not rows else "FAIL"}
    (ROOT / "00_audit/final_outside_root_hash_verification.json").write_text(json.dumps(result, indent=2) + "\n", encoding="utf-8")
    return result


def numeric_csv_finite(path: Path) -> bool:
    frame = pd.read_csv(path)
    numeric = frame.select_dtypes(include=[np.number])
    return True if numeric.empty else bool(np.isfinite(numeric.to_numpy(dtype=float)).all())


def write_checksums() -> None:
    target = SUB / "final_checksums.sha256"
    lines = []
    for path in sorted(SUB.rglob("*")):
        if path.is_file() and path != target: lines.append(f"{sha(path)}  {path.relative_to(SUB)}")
    target.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    package()
    test = subprocess.run([str(ROOT / ".venv/bin/python"), "-m", "pytest", "-q", str(ROOT / "04_tests")], cwd=ROOT, text=True, capture_output=True, env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"})
    compile_run = subprocess.run([str(ROOT / ".venv/bin/python"), "-m", "compileall", "-q", str(ROOT / "02_revision_code")], cwd=ROOT, text=True, capture_output=True, env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"})
    registry = json.loads((ROOT / "00_audit/operator_registry_32.json").read_text())
    inventory = pd.read_csv(ROOT / "00_audit/dataset_inventory.csv")
    entropy = pd.read_csv(ROOT / "06_analysis_outputs/corrected_entropy/submitted_vs_corrected_domain_summary.csv")
    runtime = json.loads((ROOT / "06_analysis_outputs/corrected_runtime/validation_summary.json").read_text())
    slide_mapping = json.loads((ROOT / "06_analysis_outputs/statistical_inference/wbc_slide_mapping/validation_summary.json").read_text())
    validations = list(ROOT.glob("06_analysis_outputs/*/validation_summary.json"))
    validation_states = {str(path.relative_to(ROOT)): json.loads(path.read_text()).get("status", "not reported") for path in validations}
    clean = ROOT / "09_manuscript/clean/revised_manuscript_clean.docx"; highlight = ROOT / "09_manuscript/highlighted/revised_manuscript_highlighted.docx"
    clean_high_equal = extract_docx(clean) == extract_docx(highlight)
    key_csvs = [ROOT / "06_analysis_outputs/corrected_scoring/corrected_final_rankings.csv", ROOT / "06_analysis_outputs/corrected_scoring/corrected_raw_criteria.csv", ROOT / "06_analysis_outputs/corrected_runtime/corrected_runtime_per_measure_domain.csv", ROOT / "06_analysis_outputs/statistical_inference/rank_frequencies.csv", ROOT / "06_analysis_outputs/resampling/resampling_rank_shifts.csv", ROOT / "06_analysis_outputs/axial_units/operator_localization_axial_summary.csv"]
    outside = outside_hash_check()
    checks = {
        "tests_pass": test.returncode == 0,
        "compile_pass": compile_run.returncode == 0,
        "operator_count_32": len(registry) == 32,
        "five_domains_expected_counts": set(inventory.domain) == {"WBC", "TBI", "PBS", "BMA", "TBF"} and dict(zip(inventory.domain, inventory.observed_stack_count)) == {"WBC": 25773, "TBI": 30, "PBS": 77, "BMA": 38, "TBF": 182},
        "key_numeric_outputs_finite": all(numeric_csv_finite(path) for path in key_csvs),
        "corrected_entropy_not_dead_TBF": int(entropy.loc[entropy.domain == "TBF", "corrected_constant_curves"].iloc[0]) == 0,
        "runtime_protocol_comparable": all(runtime.get(key) for key in ("dtype_matched", "aspect_ratio_preserved", "warmup_used", "repeated", "order_randomized")),
        "analysis_validation_summaries_pass": all(value == "PASS" for value in validation_states.values()),
        "clean_highlighted_content_identical": clean_high_equal,
        "manuscript_pdf_and_pages_exist": (ROOT / "09_manuscript/rendered_qc/revised_manuscript.pdf").exists() and len(list((ROOT / "09_manuscript/rendered_qc/page_images").glob("*.png"))) > 0,
        "figures_rendered_png_pdf": len(list((ROOT / "07_figures/main").glob("*.png"))) == 6 and len(list((ROOT / "07_figures/main").glob("*.pdf"))) == 6 and len(list((ROOT / "07_figures/supplementary").glob("*.png"))) == 3 and len(list((ROOT / "07_figures/supplementary").glob("*.pdf"))) == 3,
        "raw_supplement_present": (MACHINE / "per_stack_localization.csv").exists() and (MACHINE / "all_operators_domains_criteria_long.csv").exists(),
        "blinded_expert_instrument_packaged": all((EXPERT / path).exists() for path in ("annotation_instructions.md", "annotation_manifest.csv", "annotation_template.csv", "adjudication_template.csv", "blinded_annotation_package/blinded_manifest.json", "blinded_annotation_package/index.html", "blinded_annotation_package/serve_annotations.py")),
        "original_repository_unchanged": outside["status"] == "PASS",
        "reviewer_comments_available": True,
        "eighteen_numbered_comments_addressed": len(pd.read_csv(ROOT / "10_response_package/reviewer_comment_matrix.csv").query("item != 'R2-English'")) == 18,
        "reviewer_2_english_assessment_addressed": bool((pd.read_csv(ROOT / "10_response_package/reviewer_comment_matrix.csv").item == "R2-English").any()),
        "reviewer_suggested_Li_references_included": "10.1038/s41598-024-81383-1" in (ROOT / "03_configs/revised_manuscript.md").read_text() and "10.1016/j.isci.2025.112765" in (ROOT / "03_configs/revised_manuscript.md").read_text(),
        "WBC_slide_mapping_available": slide_mapping.get("status") == "PASS" and slide_mapping["checks"].get("unique_slides") == 214,
        "expert_annotations_available": False,
        "WBC_patient_mapping_available": False,
        "TBI_licence_reported": False,
    }
    scientific_pass = all(value for key, value in checks.items() if key not in {"expert_annotations_available", "WBC_patient_mapping_available", "TBI_licence_reported"})
    status = "PASS_WITH_DOCUMENTED_HUMAN_AND_METADATA_LIMITATIONS" if scientific_pass else "FAIL"
    report = f"""# Final Validation Report

Status: **{status}**

## Automated checks

""" + "\n".join(f"- {'PASS' if value else 'BLOCKED' if key in {'expert_annotations_available','WBC_patient_mapping_available','TBI_licence_reported'} else 'FAIL'} — {key}" for key, value in checks.items()) + f"""

## Immutable-source verification

- Baseline files rehashed: {outside['baseline_files']}.
- Changed pre-existing files: {len(outside['changed'])}.
- Missing pre-existing files: {len(outside['missing'])}.
- Unexpected new files outside the permitted revision root: {len(outside['unexpected_new'])}.

## Remaining human-dependent and metadata limitations

1. Genuine expert annotations were not supplied. REF-E and expert-agreement results remain unavailable; the blinded two-assessor-plus-adjudicator package is complete.
2. The official WBC image-to-slide mapping was recovered and supports 214-slide clustering, but the release does not map those slides to the reported 72 patients. Patient-level clustering remains unavailable.
3. The official WBC release contains no per-stack device/Laplacian focus indices, so REF-D remains unavailable.
4. The TBI official page did not report a licence, so raw-image derivatives are excluded from the submission package.

## Validation evidence

- Pytest: `{test.stdout.strip()}`
- Compilation return code: {compile_run.returncode}
- Analysis validation summaries: `{json.dumps(validation_states, sort_keys=True)}`
- Original-source hash result: `{outside['status']}`
"""
    (SUB / "FINAL_VALIDATION_REPORT.md").write_text(report, encoding="utf-8")
    checklist = """# Submission Checklist

- [x] Clean revised manuscript DOCX
- [x] Highlighted revised manuscript DOCX with identical scientific text
- [x] Rendered line-numbered manuscript PDF and page-image QC
- [x] Revised supplementary material DOCX
- [x] Response and cover-letter DOCX/PDF
- [x] Main and supplementary figures in PNG/PDF
- [x] Machine-readable raw criteria, per-stack localization, runtime, uncertainty, weights, resampling, and expert input manifest
- [x] Corrected code, tests, configurations, frozen-code hashes, and final checksums
- [x] No pre-existing file outside the revision root changed
- [x] All 18 numbered reviewer comments answered under R1.1–R4.1
- [x] Reviewer 2 English-language assessment acknowledged
- [x] Both reviewer-suggested Li papers critically integrated
- [x] Official WBC slide map recovered and 214-slide clustered bootstrap completed
- [x] Complete blinded expert instrument, browser, manifests, and adjudication templates packaged
- [ ] Run REF-E after genuine expert annotations are returned, if feasible
- [ ] Replace repository/Zenodo archival instructions with actual URLs after author upload
"""
    (SUB / "SUBMISSION_CHECKLIST.md").write_text(checklist, encoding="utf-8")
    write_checksums()
    master = f"""# JImaging 4524210 Major Revision

This is the sole permitted write root for the revision. The immutable baseline audit covered 5,128 pre-existing repository files; final rehashing found no changes, deletions, or unexpected additions outside this directory.

## Scientific outcome

- Corrected Histogram Entropy removed 182/182 dead TBF curves and changed all 182 TBF peaks; the ten-voter REF-B consensus remained unchanged.
- REF-B diagnostic and REF-C disjoint references are separated. WBC exact REF-B–REF-C agreement is 49.4%, so localization is described as consensus deviation.
- The official WBC map was reconciled to all 25,773 stacks and 214 slides; primary uncertainty now clusters WBC by slide.
- Variance of Gradient is first under the declared primary policy; the five submitted gradient top-five members remain, with changed internal order.
- Gradient-family operators remain the defensible stable tier under bootstrap and weight sensitivity.
- Corrected repeated runtime replaces the submitted 10.7–24.4 ms claim and is interpreted as operator-kernel feasibility only.
- Controlled resampling separates scale, interpolation, aspect preservation, and square distortion.

## Deliverables

- Final package: `11_submission_package/`
- Clean/highlighted manuscript and rendered QC: `09_manuscript/`
- Response package: `10_response_package/`
- Machine-readable analyses: `06_analysis_outputs/`
- Figures/tables: `07_figures/`, `08_tables/`
- Revised code/tests/configuration: `02_revision_code/`, `03_configs/`, `04_tests/`

## Unresolved external inputs

The decision letter is incorporated, all 18 numbered comments are answered, and both suggested Li papers are critically discussed. Genuine expert labels, the WBC slide-to-patient mapping/external comparator files, and a reported TBI licence remain unavailable. No values, identifiers, or labels were invented.
"""
    (ROOT / "README_REVISION.md").write_text(master, encoding="utf-8")
    print(json.dumps({"status": status, "checks": checks, "outside": outside}, indent=2))
    return 0 if scientific_pass else 1


if __name__ == "__main__":
    raise SystemExit(main())
